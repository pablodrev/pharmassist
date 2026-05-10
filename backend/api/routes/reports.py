"""Report routes."""

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, desc, func
from datetime import datetime, timedelta
from uuid import UUID
import os
import tempfile
import logging
from typing import Optional

from db import get_session
from models.schemas_db import User, Report, Drug, AIRecommendation
from auth import get_current_user, require_specialist
from api_schemas import (
    CreateReportRequest, CreateReportFromFormRequest, ReportResponse,
    ReportShortResponse, ReportListResponse, AnalysisStatusResponse,
    ChangeReportStatusRequest, SpecialistReviewRequest, FinalizeReportResponse,
    SimpleOkResponse, SeverityLevel, DateRangeFilter, ReportStatus,
    ExtractedCaseData, AIRecommendationsResponse, DrugResponse,
    ReanalyzeReportRequest
)
from services.orchestrator import AnalysisOrchestrator
from core.llm_client import LLMClient
from core.rag_engine import RAGEngine
from services.case_extraction import CaseExtractionService
import json
import requests

# Configure logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

router = APIRouter(prefix="/api/v1/reports", tags=["reports"])

# Initialize services with environment variables
def _init_llm(provider: Optional[str] = None, model: Optional[str] = None, validate_connection: bool = False):
    """Initialize LLM client from environment variables or parameters.
    
    Args:
        provider: LLM provider ('yandex' or 'ollama')
        model: Model name
        validate_connection: If True, check connection on init (only for Ollama reanalyze)
    """
    llm_provider = (provider or os.getenv("LLM_PROVIDER", "yandex")).lower()
    llm_model = model or os.getenv("LLM_MODEL", "gpt-4o-mini")
    
    logger.info(f"🔧 Initializing LLM: provider={llm_provider}, model={llm_model}")
    
    if llm_provider == "yandex":
        yandex_folder = os.getenv("YANDEX_CLOUD_FOLDER")
        yandex_key = os.getenv("YANDEX_CLOUD_API_KEY")
        
        if not yandex_folder or not yandex_key:
            error_msg = "YANDEX_CLOUD_FOLDER and YANDEX_CLOUD_API_KEY must be set"
            logger.error(f"❌ {error_msg}")
            raise ValueError(error_msg)
        
        logger.info(f"✅ Yandex Cloud configured: folder={yandex_folder[:10]}...")
        return LLMClient(
            model=llm_model,
            provider="yandex",
            yandex_api_key=yandex_key,
            yandex_folder_id=yandex_folder
        )
    else:
        ollama_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
        logger.info(f"⚙️ Ollama configured: url={ollama_url}")
        
        # Only validate connection if explicitly requested (for reanalyze)
        if validate_connection:
            try:
                response = requests.head(f"{ollama_url}/api/tags", timeout=5)
                if response.status_code != 200:
                    raise RuntimeError(f"Ollama не ответил: {response.status_code}")
                logger.info(f"✅ Ollama connection verified")
            except Exception as e:
                logger.error(f"❌ Ollama connection failed: {str(e)}")
                raise HTTPException(
                    status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                    detail=f"Ollama not available at {ollama_url}. Запустите: ollama serve"
                )
        
        return LLMClient(
            model=llm_model,
            provider="ollama",
            base_url=ollama_url
        )

_rag: Optional[RAGEngine] = None
_orchestrator: Optional[AnalysisOrchestrator] = None


def _get_rag() -> RAGEngine:
    global _rag
    if _rag is None:
        _rag = RAGEngine()
    return _rag


def _get_orchestrator() -> AnalysisOrchestrator:
    global _orchestrator
    if _orchestrator is None:
        _orchestrator = AnalysisOrchestrator(_init_llm(), _get_rag())
    return _orchestrator


@router.post("", status_code=status.HTTP_202_ACCEPTED)
async def create_report(
    request: CreateReportRequest,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Create report from raw text."""
    user_id = UUID(current_user["user_id"])
    
    logger.info(f"📝 Creating new report for user {user_id}")
    logger.debug(f"Report text length: {len(request.raw_text)} chars")
    
    # Create report
    report = Report(
        reporter_id=user_id,
        raw_text=request.raw_text,
        status="submitted"
    )
    session.add(report)
    await session.commit()
    await session.refresh(report)
    
    logger.info(f"✅ Report created with ID: {report.id}")
    
    # Run analysis in background (simplified - in production use Celery)
    try:
        logger.info(f"🔄 Starting analysis for report {report.id}...")
        analysis = _get_orchestrator().analyze(request.raw_text)
        logger.info(f"✅ Analysis completed successfully")
        
        # Convert to dict
        extracted_dict = analysis.case_extraction.model_dump() if analysis.case_extraction else None
        logger.info(f"📊 Case extraction: {bool(extracted_dict)}")
        
        # Try to match drug
        if analysis.case_extraction and analysis.case_extraction.suspect_drug:
            drug_name = analysis.case_extraction.suspect_drug.name
            logger.info(f"💊 Found drug: {drug_name}")
            stmt = select(Drug).where(Drug.name_ru == drug_name)
            result = await session.execute(stmt)
            drug = result.scalar_one_or_none()
            if not drug:
                # Create drug
                logger.info(f"➕ Creating new drug entry: {drug_name}")
                drug = Drug(name_ru=drug_name)
                session.add(drug)
                await session.commit()
                await session.refresh(drug)
            report.drug_id = drug.id
        
        # Save extracted data
        report.extracted_data = extracted_dict
        report.status = "submitted"
        
        # Save AI recommendations
        if analysis.case_extraction:
            logger.info(f"💾 Saving case_extraction recommendation")
            rec = AIRecommendation(
                report_id=report.id,
                type="case_extraction",
                ai_output=extracted_dict or {}
            )
            session.add(rec)
        
        if analysis.ime_assessment:
            logger.info(f"💾 Saving ime recommendation")
            rec = AIRecommendation(
                report_id=report.id,
                type="ime",
                ai_output=analysis.ime_assessment.model_dump()
            )
            session.add(rec)
        
        if analysis.naranjo_assessment:
            logger.info(f"💾 Saving naranjo recommendation")
            rec = AIRecommendation(
                report_id=report.id,
                type="naranjo",
                ai_output=analysis.naranjo_assessment.model_dump()
            )
            session.add(rec)
        
        if analysis.expectedness_assessment:
            logger.info(f"💾 Saving expectedness recommendation")
            rec = AIRecommendation(
                report_id=report.id,
                type="expectedness",
                ai_output=analysis.expectedness_assessment.model_dump()
            )
            session.add(rec)
        
        # Completeness
        completeness = {
            "missing_mandatory_fields": analysis.missing_mandatory_fields,
            "warnings": analysis.warnings
        }
        logger.info(f"💾 Saving completeness recommendation")
        rec = AIRecommendation(
            report_id=report.id,
            type="completeness",
            ai_output=completeness
        )
        session.add(rec)
        
        await session.commit()
        logger.info(f"✅ All recommendations saved for report {report.id}")
        
    except Exception as e:
        logger.error(f"❌ Error analyzing report {report.id}: {str(e)}", exc_info=True)
        # Still save the report, mark as failed
        rec = AIRecommendation(
            report_id=report.id,
            type="completeness",
            ai_output={"error": str(e), "missing_mandatory_fields": [], "warnings": []}
        )
        session.add(rec)
        await session.commit()
    
    return {"id": report.id, "status": report.status}


@router.post("/from-form", status_code=status.HTTP_202_ACCEPTED)
async def create_report_from_form(
    request: CreateReportFromFormRequest,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Create report from structured form."""
    user_id = UUID(current_user["user_id"])
    
    # Build narrative
    narrative = f"""
Пациент: {request.patient.name or 'Не указан'}, возраст {request.patient.age or 'не указан'}
Пол: {request.patient.sex or 'не указан'}
Вес: {request.patient.weight or 'не указан'}
Диагноз: {request.patient.diagnosis or 'не указан'}
Сопутствующие заболевания: {request.patient.comorbidities or 'нет'}

Врач: {request.doctor.name or 'Не указан'}
Специальность: {request.doctor.specialty or 'не указана'}
Организация: {request.doctor.organization or 'не указана'}
Email: {request.doctor.email or 'не указан'}

Препарат: {request.medication.trade_name or request.medication.inn or 'не указан'}
МНН: {request.medication.inn or 'не указан'}
Доза: {request.medication.dose or 'не указана'}
Путь введения: {request.medication.route or 'не указан'}
Начало: {request.medication.start_date or 'не указано'}
Окончание: {request.medication.end_date or 'не указано'}
Показание: {request.medication.indication or 'не указано'}
Производитель: {request.medication.manufacturer or 'не указан'}

Нежелательная реакция: {request.adverse_effect.description}
Дата: {request.adverse_effect.date or 'не указана'}
Тяжесть: {request.adverse_effect.severity or 'не указана'}
Серьёзная: {request.adverse_effect.is_serious or 'не указано'}
Исход: {request.adverse_effect.outcome or 'не указан'}
Оценка причинности: {request.adverse_effect.causality_assessment or 'не указана'}

Дополнительная информация: {request.additional_info.additional_info or 'нет'}
    """
    
    # Create report with raw text
    report = Report(
        reporter_id=user_id,
        raw_text=narrative,
        status="submitted"
    )
    session.add(report)
    await session.commit()
    await session.refresh(report)
    
    # Run analysis
    try:
        analysis = _get_orchestrator().analyze(narrative)
        extracted_dict = analysis.case_extraction.model_dump() if analysis.case_extraction else None
        
        # Match drug
        if analysis.case_extraction and analysis.case_extraction.suspect_drug:
            drug_name = analysis.case_extraction.suspect_drug.name
            stmt = select(Drug).where(Drug.name_ru == drug_name)
            result = await session.execute(stmt)
            drug = result.scalar_one_or_none()
            if not drug:
                drug = Drug(name_ru=drug_name)
                session.add(drug)
                await session.commit()
                await session.refresh(drug)
            report.drug_id = drug.id
        
        report.extracted_data = extracted_dict
        
        # Save recommendations
        if analysis.case_extraction:
            rec = AIRecommendation(
                report_id=report.id,
                type="case_extraction",
                ai_output=extracted_dict or {}
            )
            session.add(rec)
        
        if analysis.ime_assessment:
            rec = AIRecommendation(
                report_id=report.id,
                type="ime",
                ai_output=analysis.ime_assessment.model_dump()
            )
            session.add(rec)
        
        if analysis.naranjo_assessment:
            rec = AIRecommendation(
                report_id=report.id,
                type="naranjo",
                ai_output=analysis.naranjo_assessment.model_dump()
            )
            session.add(rec)
        
        if analysis.expectedness_assessment:
            rec = AIRecommendation(
                report_id=report.id,
                type="expectedness",
                ai_output=analysis.expectedness_assessment.model_dump()
            )
            session.add(rec)
        
        completeness = {
            "missing_mandatory_fields": analysis.missing_mandatory_fields,
            "warnings": analysis.warnings
        }
        rec = AIRecommendation(
            report_id=report.id,
            type="completeness",
            ai_output=completeness
        )
        session.add(rec)
        
        await session.commit()
        
    except Exception as e:
        print(f"Error analyzing report: {e}")
        rec = AIRecommendation(
            report_id=report.id,
            type="completeness",
            ai_output={"error": str(e), "missing_mandatory_fields": [], "warnings": []}
        )
        session.add(rec)
        await session.commit()
    
    return {"id": report.id, "status": report.status}


@router.post("/extract-from-file", response_model=ExtractedCaseData)
async def extract_from_file(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    """Extract data from PDF/DOCX without creating report."""
    # Определяем тип файла по MIME или расширению имени (Windows может слать application/octet-stream для .docx)
    filename_lower = (file.filename or "").lower()
    _PDF_MIME = "application/pdf"
    _DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if file.content_type == _PDF_MIME or filename_lower.endswith(".pdf"):
        effective_type = "pdf"
    elif file.content_type == _DOCX_MIME or filename_lower.endswith(".docx"):
        effective_type = "docx"
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX files are supported"
        )

    if file.size > 10 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File size must be less than 10 MB"
        )

    logger.info(
        "[extract-from-file] filename=%r content_type=%r effective_type=%r size=%s",
        file.filename, file.content_type, effective_type, file.size,
    )

    # Save file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    logger.info("[extract-from-file] saved to tmp=%r (%d bytes)", tmp_path, len(content))

    try:
        # Извлечение текста из файла
        text = ""
        if effective_type == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(tmp_path)
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages_text)
            logger.info("[extract-from-file] PDF: %d pages, total chars=%d", len(pages_text), len(text))
        else:  # effective_type == "docx"
            from docx import Document
            doc = Document(tmp_path)
            parts: list[str] = []
            para_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
                    para_count += 1
            table_count = 0
            cell_count = 0
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(" | ".join(row_cells))
                        cell_count += len(row_cells)
                table_count += 1
            text = "\n".join(parts)
            logger.info(
                "[extract-from-file] DOCX: paragraphs=%d tables=%d cells=%d total_chars=%d",
                para_count, table_count, cell_count, len(text),
            )

        logger.info(
            "[extract-from-file] extracted text preview: %r",
            text[:600],
        )

        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Не удалось извлечь текст из файла. "
                       "Убедитесь, что документ содержит текстовые данные (не только изображения).",
            )
        
        # Run case extraction only
        case_svc = CaseExtractionService(_init_llm())
        extraction = case_svc.extract(text)
        
        return ExtractedCaseData(
            patient=extraction.patient.model_dump() if extraction.patient else None,
            reporter=extraction.reporter.model_dump() if extraction.reporter else None,
            adverse_reaction=extraction.adverse_reaction.model_dump() if extraction.adverse_reaction else None,
            suspect_drug=extraction.suspect_drug.model_dump() if extraction.suspect_drug else None,
            concomitant_drugs=[d.model_dump() for d in extraction.concomitant_drugs],
            case_narrative=extraction.case_narrative,
            raw_text=text,  # передаём сырой текст фронтенду как fallback
        )
    finally:
        os.unlink(tmp_path)


@router.get("", response_model=ReportListResponse)
async def list_reports(
    status_filter: Optional[str] = Query(None, alias="status"),
    search: Optional[str] = Query(None),
    severity: Optional[SeverityLevel] = Query(None),
    date_range: Optional[DateRangeFilter] = Query(None),
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """List reports with filters."""
    user_id = UUID(current_user["user_id"])
    role = current_user["role"]
    
    # Build query
    query = select(Report)
    
    # Filter by user role
    if role == "reporter":
        query = query.where(Report.reporter_id == user_id)
    
    # Filter by status
    if status_filter:
        query = query.where(Report.status == status_filter)
    
    # Filter by date range
    if date_range:
        now = datetime.utcnow()
        if date_range == DateRangeFilter.TODAY:
            start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif date_range == DateRangeFilter.WEEK:
            start_date = now - timedelta(days=7)
        elif date_range == DateRangeFilter.MONTH:
            start_date = now - timedelta(days=30)
        else:
            start_date = None
        if start_date:
            query = query.where(Report.created_at >= start_date)
    
    # Get total count
    count_stmt = select(func.count(Report.id)).select_from(Report).where(
        query.whereclause if query.whereclause is not None else True
    )
    count_result = await session.execute(count_stmt)
    total = count_result.scalar() or 0
    
    # Paginate
    query = query.order_by(desc(Report.created_at))
    query = query.offset((page - 1) * limit).limit(limit)
    
    result = await session.execute(query)
    reports = result.scalars().all()
    
    # Build responses
    items = []
    for report in reports:
        # Get AI recommendations
        stmt = select(AIRecommendation).where(AIRecommendation.report_id == report.id)
        recs_result = await session.execute(stmt)
        recs = recs_result.scalars().all()
        
        ime_rec = next((r for r in recs if r.type == "ime"), None)
        
        # Extract drug name
        drug_name = None
        if report.drug_id:
            stmt = select(Drug).where(Drug.id == report.drug_id)
            drug_result = await session.execute(stmt)
            drug = drug_result.scalar_one_or_none()
            if drug:
                drug_name = drug.name_ru
        
        # Extract adverse reaction and reporter
        adverse_reaction = None
        reporter_name = None
        if report.extracted_data:
            data = report.extracted_data
            if isinstance(data, str):
                data = json.loads(data)
            if data.get("adverse_reaction"):
                adverse_reaction = data["adverse_reaction"].get("description")
            if data.get("reporter"):
                reporter_name = data["reporter"].get("name")
        
        item = ReportShortResponse(
            id=report.id,
            status=report.status,
            reporter_id=report.reporter_id,
            drug_name=drug_name,
            adverse_reaction=adverse_reaction,
            reporter_name=reporter_name,
            is_clinically_significant=ime_rec.ai_output.get("is_clinically_significant") if ime_rec else None,
            analysis_status="ready" if ime_rec else "pending",
            created_at=report.created_at,
            severity=None  # TODO: extract from data
        )
        items.append(item)
    
    return ReportListResponse(items=items, total=total, page=page)


@router.get("/{report_id}", response_model=ReportResponse)
async def get_report(
    report_id: UUID,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Get report details."""
    user_id = UUID(current_user["user_id"])
    role = current_user["role"]
    
    stmt = select(Report).where(Report.id == report_id)
    result = await session.execute(stmt)
    report = result.scalar_one_or_none()
    
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    # Check access
    if role == "reporter" and report.reporter_id != user_id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")
    
    # Get drug
    drug = None
    if report.drug_id:
        stmt = select(Drug).where(Drug.id == report.drug_id)
        result = await session.execute(stmt)
        drug_obj = result.scalar_one_or_none()
        if drug_obj:
            drug = DrugResponse(
                id=drug_obj.id,
                name_ru=drug_obj.name_ru,
                name_en=drug_obj.name_en,
                atc_code=drug_obj.atc_code
            )
    
    # Get AI recommendations
    stmt = select(AIRecommendation).where(AIRecommendation.report_id == report.id)
    result = await session.execute(stmt)
    recs = result.scalars().all()
    
    ai_recs = {}
    for rec in recs:
        ai_recs[rec.type] = rec.ai_output
    
    analysis_status = "ready" if ai_recs else "pending"
    
    # Build extracted data
    extracted_data = None
    if report.extracted_data:
        data = report.extracted_data
        if isinstance(data, str):
            data = json.loads(data)
        extracted_data = ExtractedCaseData(**data)
    
    ai_recommendations = AIRecommendationsResponse(
        analysis_status=analysis_status,
        case_extraction=extracted_data,
        ime=ai_recs.get("ime"),
        naranjo=ai_recs.get("naranjo"),
        expectedness=ai_recs.get("expectedness"),
        completeness=ai_recs.get("completeness")
    )
    
    return ReportResponse(
        id=report.id,
        status=report.status,
        reporter_id=report.reporter_id,
        raw_text=report.raw_text,
        created_at=report.created_at,
        drug=drug,
        extracted_data=extracted_data,
        ai_recommendations=ai_recommendations
    )


@router.get("/{report_id}/analysis-status", response_model=AnalysisStatusResponse)
async def get_analysis_status(
    report_id: UUID,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session)
):
    """Get AI analysis status."""
    stmt = select(Report).where(Report.id == report_id)
    result = await session.execute(stmt)
    report = result.scalar_one_or_none()
    
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    # Get completeness recommendation
    stmt = select(AIRecommendation).where(
        and_(AIRecommendation.report_id == report.id, AIRecommendation.type == "completeness")
    )
    result = await session.execute(stmt)
    rec = result.scalar_one_or_none()
    
    if not rec:
        return AnalysisStatusResponse(analysis_status="pending")
    
    error = rec.ai_output.get("error") if rec.ai_output else None
    status_val = "failed" if error else "ready"
    
    return AnalysisStatusResponse(
        analysis_status=status_val,
        error=error
    )


@router.patch("/{report_id}/status", response_model=SimpleOkResponse)
async def change_report_status(
    report_id: UUID,
    request: ChangeReportStatusRequest,
    current_user: dict = Depends(require_specialist),
    session: AsyncSession = Depends(get_session)
):
    """Change report status. Only specialists."""
    stmt = select(Report).where(Report.id == report_id)
    result = await session.execute(stmt)
    report = result.scalar_one_or_none()
    
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    # Validate transition
    if request.status not in ["clarification", "analysis"]:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid status")
    
    report.status = request.status
    session.add(report)
    await session.commit()
    
    return SimpleOkResponse(ok=True)


@router.patch("/{report_id}/specialist-review", response_model=SimpleOkResponse)
async def specialist_review(
    report_id: UUID,
    request: SpecialistReviewRequest,
    current_user: dict = Depends(require_specialist),
    session: AsyncSession = Depends(get_session)
):
    """Save specialist overrides."""
    stmt = select(Report).where(Report.id == report_id)
    result = await session.execute(stmt)
    report = result.scalar_one_or_none()
    
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    # Get AI recommendations
    stmt = select(AIRecommendation).where(AIRecommendation.report_id == report.id)
    result = await session.execute(stmt)
    recs = result.scalars().all()
    
    # Update verdicts
    rec_map = {r.type: r for r in recs}
    
    if request.causality and "causality" in rec_map:
        rec_map["causality"].specialist_verdict = request.causality.get("verdict")
        rec_map["causality"].specialist_comment = request.causality.get("comment")
    
    if request.ime and "ime" in rec_map:
        rec_map["ime"].specialist_verdict = request.ime.get("verdict")
        rec_map["ime"].specialist_comment = request.ime.get("comment")
    
    if request.naranjo and "naranjo" in rec_map:
        rec_map["naranjo"].specialist_verdict = request.naranjo.get("verdict")
        rec_map["naranjo"].specialist_comment = request.naranjo.get("comment")
    
    if request.expectedness and "expectedness" in rec_map:
        rec_map["expectedness"].specialist_verdict = request.expectedness.get("verdict")
        rec_map["expectedness"].specialist_comment = request.expectedness.get("comment")
    
    if request.completeness and "completeness" in rec_map:
        rec_map["completeness"].specialist_verdict = request.completeness.get("verdict")
        rec_map["completeness"].specialist_comment = request.completeness.get("comment")
    
    for rec in recs:
        session.add(rec)
    
    await session.commit()
    
    return SimpleOkResponse(ok=True)


@router.post("/{report_id}/finalize", response_model=FinalizeReportResponse)
async def finalize_report(
    report_id: UUID,
    current_user: dict = Depends(require_specialist),
    session: AsyncSession = Depends(get_session)
):
    """Finalize report. Only specialists."""
    stmt = select(Report).where(Report.id == report_id)
    result = await session.execute(stmt)
    report = result.scalar_one_or_none()
    
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    report.status = "finalized"
    session.add(report)
    await session.commit()
    
    return FinalizeReportResponse(ok=True, finalized_at=datetime.utcnow())


@router.post("/{report_id}/reanalyze", response_model=SimpleOkResponse)
async def reanalyze_report(
    report_id: UUID,
    request: ReanalyzeReportRequest,
    current_user: dict = Depends(require_specialist),
    session: AsyncSession = Depends(get_session)
):
    """
    Re-analyze report with a different LLM provider/model.
    Only specialists can use this endpoint.
    This overwrites previous analysis results (Variant A - simple).
    """
    logger.info(f"🔄 Reanalyze request for report {report_id}")
    logger.info(f"   Requested: provider={request.llm_provider}, model={request.llm_model}")
    
    # Validate provider
    if request.llm_provider.lower() not in ["yandex", "ollama"]:
        logger.warning(f"❌ Invalid provider: {request.llm_provider}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Provider must be 'yandex' or 'ollama'"
        )
    
    # Get report
    stmt = select(Report).where(Report.id == report_id)
    result = await session.execute(stmt)
    report = result.scalar_one_or_none()
    
    if not report:
        logger.warning(f"❌ Report {report_id} not found")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Report not found"
        )
    
    logger.info(f"✅ Found report {report_id}, raw_text length: {len(report.raw_text)} chars")
    
    try:
        # Initialize new LLM client with requested parameters
        logger.info(f"🔧 Creating new LLM client...")
        
        # For Ollama, validate connection first
        if request.llm_provider.lower() == "ollama":
            logger.info(f"🔍 Checking Ollama availability...")
            new_llm = _init_llm(
                provider=request.llm_provider.lower(),
                model=request.llm_model,
                validate_connection=True
            )
        else:
            new_llm = _init_llm(
                provider=request.llm_provider.lower(),
                model=request.llm_model,
                validate_connection=False
            )
        
        logger.info(f"✅ New LLM client initialized")
        
        # Create new orchestrator with new LLM
        new_orchestrator = AnalysisOrchestrator(new_llm, _get_rag())
        logger.info(f"🔄 Starting re-analysis with {request.llm_provider}/{request.llm_model}...")
        
        # Run analysis
        analysis = new_orchestrator.analyze(report.raw_text)
        logger.info(f"✅ Re-analysis completed successfully")
        
        # Delete existing recommendations (Variant A - overwrite)
        logger.info(f"🗑️ Deleting previous recommendations...")
        stmt = select(AIRecommendation).where(AIRecommendation.report_id == report.id)
        result = await session.execute(stmt)
        existing_recs = result.scalars().all()
        for rec in existing_recs:
            await session.delete(rec)
        logger.info(f"✅ Deleted {len(existing_recs)} previous recommendations")
        
        # Convert to dict
        extracted_dict = analysis.case_extraction.model_dump() if analysis.case_extraction else None
        
        # Save extracted data
        report.extracted_data = extracted_dict
        logger.info(f"📊 Updated extracted data")
        
        # Save new AI recommendations
        if analysis.case_extraction:
            logger.info(f"💾 Saving case_extraction from {request.llm_provider}")
            rec = AIRecommendation(
                report_id=report.id,
                type="case_extraction",
                ai_output=extracted_dict or {}
            )
            session.add(rec)
        
        if analysis.ime_assessment:
            logger.info(f"💾 Saving ime from {request.llm_provider}")
            rec = AIRecommendation(
                report_id=report.id,
                type="ime",
                ai_output=analysis.ime_assessment.model_dump()
            )
            session.add(rec)
        
        if analysis.naranjo_assessment:
            logger.info(f"💾 Saving naranjo from {request.llm_provider}")
            rec = AIRecommendation(
                report_id=report.id,
                type="naranjo",
                ai_output=analysis.naranjo_assessment.model_dump()
            )
            session.add(rec)
        
        if analysis.expectedness_assessment:
            logger.info(f"💾 Saving expectedness from {request.llm_provider}")
            rec = AIRecommendation(
                report_id=report.id,
                type="expectedness",
                ai_output=analysis.expectedness_assessment.model_dump()
            )
            session.add(rec)
        
        # Completeness
        completeness = {
            "missing_mandatory_fields": analysis.missing_mandatory_fields,
            "warnings": analysis.warnings,
            "reanalyzed_with": f"{request.llm_provider}/{request.llm_model}"
        }
        logger.info(f"💾 Saving completeness from {request.llm_provider}")
        rec = AIRecommendation(
            report_id=report.id,
            type="completeness",
            ai_output=completeness
        )
        session.add(rec)
        
        await session.commit()
        logger.info(f"✅ All new recommendations saved for report {report_id}")
        logger.info(f"✅ Re-analysis completed successfully with {request.llm_provider}/{request.llm_model}")
        
        return SimpleOkResponse(ok=True)
        
    except Exception as e:
        logger.error(f"❌ Error during re-analysis of {report_id}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Re-analysis failed: {str(e)}"
        )
